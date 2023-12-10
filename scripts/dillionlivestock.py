from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
import json
import re
import requests
from io import BytesIO
from docx.oxml.ns import qn
from datetime import datetime
import sys
sys.path.append('../helpers')
from helpers.s3 import store_data

# Function to convert date format from mm/dd/YYYY to YYYY-mm-dd
def convert_date_format(date_str):
    return datetime.strptime(date_str, '%m/%d/%Y').strftime('%Y-%m-%d')

# Function to extract the first date from the first page of a DOCX document
def extract_first_date(doc):
    first_page_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    date_match = re.search(r'\d{1,2}/\d{1,2}/\d{4}', first_page_text)
    return convert_date_format(date_match.group()) if date_match else None

def download_document(url):
    """Download a document from the given URL using requests."""
    response = requests.get(url, verify=False)
    if response.status_code == 200:
        return BytesIO(response.content)
    else:
        raise Exception(f"Failed to download the document: {response.status_code}")

def parse_document_tables(doc_stream):
    """Parse the tables from the DOCX document starting from the second page."""
    doc = Document(doc_stream)
    date = extract_first_date(doc)
    parsed_data = []
    found_second_page = False

    for table in doc.tables:
        if table.cell(0,0).text.strip() == "Name":
            found_second_page = True
        
        if found_second_page:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            for row in table.rows[1:]:
                row_data = {'Date': date,
                            'Auction': 'Dillon'}
                for idx, cell in enumerate(row.cells):
                    if cell._element.get(qn('w:vMerge')) is not None and \
                       cell._element.get(qn('w:vMerge')).get(qn('w:val')) == 'continue':
                        text = ''
                    else:
                        text = cell.text.strip()
                    header = headers[idx] if idx < len(headers) else f'Extra_{idx}'
                    row_data[header] = text
                parsed_data.append(row_data)
    
    return parsed_data, date

def run_scrape(driver):
    # URL of the page where the documents are listed
    url = 'https://dillonlivestockauction.com/market-reports/'

    # Get the first three document links
    driver.get(url)
    WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.CSS_SELECTOR, "a.wp-block-file__button"))
    )
    link = driver.find_element(By.CSS_SELECTOR, "a.wp-block-file__button")
    document_url = link.get_attribute('href')

    # Download and parse each document
    all_data = []
    try:
        doc_stream = download_document(document_url)
        parsed_data, date = parse_document_tables(doc_stream)
        all_data.extend(parsed_data)
    except Exception as e:
        print(f"An error occurred: {e}")

    # Filter the data to only include rows where all columns are populated
    filtered_data = [row for row in all_data if all(row.values())]
    store_data(date, filtered_data, "cattleiq/dillionauction")
    
